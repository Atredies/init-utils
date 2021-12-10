# Simple interface for generating DOCX files with text
# This is not fully working, just an example on what different components do

# Import DOCX Generation Libs
from docx import Document
from docx.text.paragraph import Paragraph
import os

# Import for Templates selection
from string import Template

# Import GUI Libs
import PySimpleGUI as sg

# STATIC VARIABLES:
document = Document()

# Default Save location:
# save_dir = os.environ['HOME']     # Works only on Linux

save_dir = os.path.expanduser('~')  # Works on Windows and Linux

# TITLE CONFIGURATION:
title1 = "Hello World, I am title 1"
title2 = "Hello World, I am title 2"
title3 = "Hello World, I am title 3"
title4 = "Hello World, I am title 4"

title_list = [
    title1, 
    title2, 
    title3, 
    title4
    ]

# PARAGRAPHS CONFIGURATION:
paragraph1_p1 = "The Client's name is: "
paragraph1_p2 = ". Well done"
paragraph3 = "Cool"
paragraph4 = paragraph1_p2 + "\n" + paragraph1_p2

# TEMPLATE CONFIGURATION:
template1 = "hello1"
template2 = "hello2"
template3 = "hello3"
template4 = "hello4"

template_list = [
    template1,
    template2,
    template3,
    template4
]

# GUI Layout:
layout = [
# Title Selection
[sg.Text("Choose Title:", font="Courier 12", text_color="green", background_color="black")],
[sg.Radio(title_list[0], "RADIO1", default=True, key="-TITLE1-")],
[sg.Radio(title_list[1], "RADIO1", default=False, key="-TITLE2-")],
[sg.Radio(title_list[2], "RADIO1", default=False, key="-TITLE3-")],
[sg.Radio(title_list[3], "RADIO1", default=False, key="-TITLE4-")],
[sg.Text("\n")],

# Template Selection Section
[sg.Text("Select Template:", font="Courier 12", text_color="green", background_color="black")],
[sg.Radio(template_list[0], "RADIO2", default=True, key="-TEMPLATE1-")],
[sg.Radio(template_list[1], "RADIO2", default=False, key="-TEMPLATE2-")],
[sg.Radio(template_list[2], "RADIO2", default=False, key="-TEMPLATE3-")],
[sg.Radio(template_list[3], "RADIO2", default=False, key="-TEMPLATE4-")],

# Manual Input Section:
[sg.Text("Detail Completion Section:", font="Courier 12", text_color="green", background_color="black")],
[sg.Text("Enter Full Name of Client:"), sg.Input("", key="-CLIENTNAME-"),],
[sg.Text("Enter Client Details:"), sg.Input("", key="-CLIENTDETAILS-")],
[sg.Text("\n")],
[sg.Text("Choose Report Name and Location:", font="Courier 12", text_color="green", background_color="black")],
[sg.Text("Select Filename:", size=(15,1 )), sg.Input("", key="-SAVEAS-"), sg.FileSaveAs()],

# Save / Cancel Section
[sg.Button("Save"), sg.Button("Cancel")],
]


# GUI Create Window:
window = sg.Window("Template Generator", layout, margins=(200, 200))

# GUI LOGIC:
def gui_logic():
    def report_save_func():
        if event == "Save":
            if values["-SAVEAS-"]:
                filename = values["-SAVEAS-"]
                document.save(filename + ".docx")
            else:
                sg.popup_error("An Error has occured! Please select a save location")
        elif event == "Cancel":
            window.close()
        else:
            print("Was not able to determine button")
            window.close()

    def client_name_func():
        if values["-CLIENTNAME-"]:
            paragraph_text = paragraph1_p1 + values["-CLIENTNAME-"] + paragraph1_p2
            p = document.add_paragraph(paragraph_text + "\n")
            p.add_run(paragraph_text + "\n").bold = True
            p.add_run(paragraph_text + "\n").italic = True
        else:
            print("Not able to get the value from Client Name field")

    def client_details_func():
        if values["-CLIENTDETAILS-"]:
            paragraph_text = paragraph1_p1 + values["-CLIENTDETAILS-"] + paragraph1_p2
            p = document.add_paragraph(paragraph_text + "\n")
            p.add_run(paragraph_text + "\n").bold = True
            p.add_run(paragraph_text + "\n").italic = True
        else:
            print("Not able to get the value from Details field")
    
    def template_select_func():
        if values["-TEMPLATE1-"]:
            template_text =  template_list[0]
            p = document.add_paragraph(template_text)
        elif values["-TEMPLATE2-"]:
            template_text = template_list[1]
            p = document.add_paragraph(template_text)
        elif values["-TEMPLATE3-"]:
            template_text = template_list[2]
            p = document.add_paragraph(template_text)
        elif values["-TEMPLATE4-"]:
            template_text = template_list[3]
            p = document.add_paragraph(template_text)
        else:
            print("Was unable to determine what template to choose")




##### BUTTON FUNCTIONS #####
    while True:
        event, values = window.read()
        if event == sg.WIN_CLOSED or event=="Exit":
            break
        elif values["-TITLE1-"]:
            title = title_list[0]
            document.add_heading(title, 0)
            template_select_func()
            client_name_func()
            client_details_func()
            report_save_func()
        elif values["-TITLE2-"]:
            title = title_list[1]
            document.add_heading(title, 0)
            template_select_func()
            client_name_func()
            client_details_func()
            report_save_func()
        elif values["-TITLE3-"]:
            title = title_list[2]
            document.add_heading(title, 0)
            template_select_func()
            client_name_func()
            client_details_func()
            report_save_func()
        elif values["-TITLE4-"]:
            title = title_list[3]
            document.add_heading(title, 0)
            template_select_func()
            client_name_func()
            client_details_func()

    window.close

gui_logic()